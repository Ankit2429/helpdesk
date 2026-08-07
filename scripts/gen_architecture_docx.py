"""Generate Campus_Helpdesk_Complete_Architecture.docx.

Consumes the per-subsystem JSON reports written to .architecture_scan/ by the
architecture-scan workflow, plus a hardcoded overview of the project, and emits
a complete, structured Word document covering every folder, every file, the
architecture, the glossary, configuration, risks and data artifacts.

Usage:
    .venv/Scripts/python.exe scripts/gen_architecture_docx.py [output.docx]
"""

from __future__ import annotations

import datetime
import glob
import json
import os
import re
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SCAN_DIR = os.path.join(REPO_ROOT, ".architecture_scan")

ACCENT = RGBColor(0x1F, 0x4E, 0x79)   # dark blue
ACCENT_LIGHT = "DEEAF6"                # light blue for table headers
GRAY = RGBColor(0x59, 0x59, 0x59)
MONO = "Consolas"


# ---------------------------------------------------------------- helpers

def set_run(run, size=10, bold=False, italic=False, mono=False, color=None):
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    if mono:
        run.font.name = MONO
        run.font.size = Pt(max(size - 0.5, 7.5))
    return run


def para(doc, text, size=10.5, bold=False, italic=False, style=None,
         align=None, space_after=6, color=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        set_run(p.add_run(text), size=size, bold=bold, italic=italic, color=color)
    return p


def add_rich(doc, text, size=10.5):
    """Add a paragraph supporting **bold** and `code` inline markers."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    for token in re.split(r"(\*\*.*?\*\*|`.*?`)", text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            set_run(p.add_run(token[2:-2]), size=size, bold=True)
        elif token.startswith("`") and token.endswith("`"):
            set_run(p.add_run(token[1:-1]), size=size, mono=True)
        else:
            set_run(p.add_run(token), size=size)
    return p


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def cell_text(cell, text, size=9, bold=False, mono=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    if isinstance(text, str) and text:
        lines = text.split("\n")
        for i, ln in enumerate(lines):
            if i:
                p = cell.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
            set_run(p.add_run(ln), size=size, bold=bold, mono=mono, color=color)


def mono_para(doc, text, size=8.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    for line in text.split("\n"):
        run = p.add_run(line)
        run.font.name = MONO
        run.font.size = Pt(size)
        p.add_run("\n")
    return p


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fld = OxmlElement("w:fldChar"); fld.set(qn("w:fldCharType"), "begin"); fld.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = "Table of contents — right-click here and choose \u201cUpdate Field\u201d to populate."
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    for el in (fld, instr, sep, t, end):
        run._r.append(el)


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_field_table(doc, headers, rows, widths=None, size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        cell_text(hdr[i], h, size=size, bold=True)
        shade_cell(hdr[i], ACCENT_LIGHT)
    for r in rows:
        cells = table.add_row().cells
        for i, val in enumerate(r):
            cell_text(cells[i], val, size=size, mono=(i == 0 and headers[0] == "File"))
    if widths:
        for i, w in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table


def h1(doc, text):
    h = doc.add_heading(text, level=1)
    for r in h.runs:
        r.font.color.rgb = ACCENT
    return h


def h2(doc, text):
    h = doc.add_heading(text, level=2)
    for r in h.runs:
        r.font.color.rgb = ACCENT
    return h


def h3(doc, text):
    h = doc.add_heading(text, level=3)
    for r in h.runs:
        r.font.color.rgb = ACCENT
    return h


def slugify(s):
    return re.sub(r"[^0-9A-Za-z]+", "_", s).strip("_").lower()


# ------------------------------------------------------------- report loading

def load_reports():
    """Load scan reports keyed by their unique filename stem (e.g. '01_wiring').

    The folder field is NOT unique across reports (e.g. three 'scripts/' reports),
    so the report filename is the only reliable key.
    """
    reports = {}
    if not os.path.isdir(SCAN_DIR):
        return reports
    for fp in sorted(glob.glob(os.path.join(SCAN_DIR, "*.json"))):
        stem = os.path.splitext(os.path.basename(fp))[0]
        try:
            with open(fp, encoding="utf-8") as f:
                data = json.load(f)
            if isinstance(data, dict):
                reports[stem] = data
        except Exception as exc:  # noqa: BLE001 - one bad report shouldn't kill the run
            print(f"  ! failed to parse {os.path.basename(fp)}: {exc}")
    return reports


def ordered_reports(reports):
    """Reports in document order (numeric filename prefixes keep the workflow order)."""
    return [(stem, reports[stem]) for stem in sorted(reports)]


def collect_terms(reports):
    """Aggregate + dedupe all per-file glossary terms across reports."""
    terms = {}
    for rep in reports.values():
        for f in rep.get("files", []):
            for t in f.get("terms", []):
                term = str(t.get("term", "")).strip()
                if not term:
                    continue
                key = term.lower()
                if key not in terms:
                    terms[key] = {"term": term, "definition": str(t.get("definition", "")).strip()}
    return sorted(terms.values(), key=lambda x: x["term"].lower())


def collect_config(reports):
    keys = {}
    for rep in reports.values():
        for f in rep.get("files", []):
            for c in f.get("config", []):
                c = str(c).strip()
                if c:
                    keys[c] = None
    return sorted(keys)


def all_file_rows(reports):
    """Every file recorded by a report -> (path, subsystem, purpose)."""
    rows = []
    for rep in reports.values():
        for f in rep.get("files", []):
            path = str(f.get("path", ""))
            if path:
                rows.append((path, rep.get("title", rep.get("folder", "")), str(f.get("purpose", "")).replace("\n", " ")))
    return rows


# ---------------------------------------------------------------- sections

def section_overview(doc):
    h1(doc, "1. Executive Summary")
    para(doc, "The Campus Helpdesk is an offline-first RAG (Retrieval-Augmented Generation) chatbot "
              "built for a college campus. It answers questions about academics, departments, faculty, "
              "facilities, fees, hostels, placement and navigation using a local knowledge base. "
              "Retrieval is hybrid (BM25 lexical + FAISS vector search fused with Reciprocal Rank Fusion), "
              "and generation runs on a local Ollama LLM so the system works without internet.")
    add_rich(doc, "There are **two parallel code stacks** that do not import each other:")
    for line in [
        "**`src/campus_helpdesk/`** — the layered, production FastAPI/RAG stack (HTTP API, application layer, domain, infrastructure, services, analytics). This is the primary surface.",
        "**Robot stack** (`interaction/`, `runtime/`, `presentation/`, `services/` voice & vision) — a separate Raspberry Pi deployment for a physical kiosk robot. Not wired into the API.",
        "**Legacy / scratch code** at the repo root (`assistant_loop.py`, `helpdesk_gui.py`, `chat.py`, `config.py`, root `services/`, `conversation/`, analysis/diagnostic scripts) and in `evaluation/`, `scripts/` — see Section 9.",
    ]:
        para(doc, line, size=10.5, space_after=4)
    add_rich(doc, "The HTTP API is the primary production surface; the robot stack is deployed separately on a "
                  "Raspberry Pi (STT/TTS/vision). The knowledge base is built by scraping the college website "
                  "into canonical Markdown (`data/canonical_markdown/`, `bvb_scraped_data/`), chunked and indexed "
                  "into FAISS + BM25. **~39,000 lines of Python** across ~270 files; the `src/` tree alone is "
                  "122 modules.")


def section_architecture(doc):
    h1(doc, "2. High-Level Architecture")
    para(doc, "The system follows a layered architecture inside `src/campus_helpdesk/`. Arrows point from the "
              "entry surface down through the layers to data and external services.", space_after=10)
    mono_para(doc, (
        "                    ┌───────────────────────────────────────────────┐\n"
        "   HTTP / Web UI   │              HTTP API (FastAPI)               │\n"
        "     (POST /chat)  │  api/routes/chat · api/routes/system          │\n"
        "                    │  api/schemas · api/dependencies (DI)          │\n"
        "                    ├───────────────────────────────────────────────┤\n"
        "                    │            Application (use cases)           │\n"
        "                    │  RAGChatService · RAGPipeline · QueryRewriter │\n"
        "                    │  SessionManager · ConversationManager · ...  │\n"
        "                    ├───────────────────────────────────────────────┤\n"
        "                    │        Domain (entities, no I/O)             │\n"
        "                    │  KnowledgeDocument · SearchResult ·          │\n"
        "                    │  ChatMessage · ConversationMemory            │\n"
        "                    ├───────────────────────────────────────────────┤\n"
        "                    │        Infrastructure (adapters)             │\n"
        "                    │  RAG: BM25+FAISS→RRF→cross-encoder rerank    │\n"
        "                    │  LLM: Ollama (local) / cloud · Knowledge     │\n"
        "                    │  pipeline · Vision · Evaluation              │\n"
        "                    ├───────────────────────────────────────────────┤\n"
        "                    │     Services (cross-cutting) · Analytics     │\n"
        "                    │  sanitizer · citation validator · language   │\n"
        "                    │  detector · answerability · metrics (dormant)│\n"
        "                    └───────────────────────────┬───────────────────┘\n"
        "                                                │\n"
        "                     ┌──────────────┬────────────┼────────────┬────────────────┐\n"
        "                     ▼              ▼            ▼            ▼                ▼\n"
        "              FAISS index      BM25 store   canonical    Ollama LLM      sentence-\n"
        "              (vectors)      (lexical n-    markdown   (local, qwen    transformer\n"
        "                                 grams)    knowledge   / qwen2.5:3b)    embeddings\n"
        "                                                     ────────────────────────\n"
        "                              ROBOT STACK (Raspberry Pi, separate deploy)\n"
        "        interaction (event bus / FSM) → runtime/system_runtime → presentation (PySide6)\n"
        "        services: STT · TTS · VAD · wake word · vision/camera · inference adapter"
    ))
    h2(doc, "2.1 The /chat request flow")
    flow = [
        ("1", "Client calls `POST /chat` → `api/routes/chat.py` → `RAGChatService.respond(message, session_id)`."),
        ("2", "`language_detector` checks the language; if not English, the query is routed for translation."),
        ("3", "`QueryRewriter` rewrites the query using the conversation history (expects `Sequence[ChatMessage]`, not a string)."),
        ("4", "`RAGPipeline.search` → `HybridRetriever` runs BM25 + FAISS, fuses scores with RRF, dedupes and returns top-k."),
        ("5", "`CrossEncoderReranker` re-ranks candidates; `ConfidenceEngine` scores them; docs are filtered by a distance threshold."),
        ("6", "`PromptContextBuilder` formats context (≈3000 char budget); a system prompt + history + context + question is assembled."),
        ("7", "`OllamaLLMService.generate` produces the reply (single `user` role message; no separate system role)."),
        ("8", "`CitationValidator` strips fabricated `[n]` citations/URLs; the reply is stored in the session's `ConversationMemory`."),
    ]
    for num, txt in flow:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        set_run(p.add_run(f"Step {num}. "), size=10.5, bold=True)
        set_run(p.add_run(txt), size=10.5)


def section_config(doc, reports):
    h1(doc, "3. Configuration & Settings")
    para(doc, "Settings are defined in `src/campus_helpdesk/config/settings.py` (Pydantic `Settings`). "
              "Sources load in order and later sources override earlier ones:")
    for line in [
        "1. `config.yaml` at the repo root (nested `app:` / `retrieval:` / `ollama:` / `embedding:` sections), flattened into settings.",
        "2. Environment variables / `.env` file (e.g. `OLLAMA_MODEL`, `RAG_SEARCH_LIMIT`).",
    ]:
        para(doc, line, size=10.5, space_after=4)
    para(doc, "Important quirk: `config.yaml` is applied through a `model_validator(mode=\"before\")` that "
              "**overwrites** same-named environment keys — `config.yaml` wins over `.env` for overlapping "
              "settings (e.g. `retrieval.top_k` → `rag_search_limit`, `ollama.model` → `ollama_model`). "
              "Unknown keys are ignored (`extra=\"ignore\"`).", italic=True)
    para(doc, "Notable settings:", bold=True, space_after=4)
    for line in [
        "`FAISS_ALLOW_DANGEROUS_DESERIALIZATION` — gates index loading (off by default; the pipeline runs degraded without a loaded index).",
        "`EMBEDDING_LOCAL_FILES_ONLY=true` — keeps the embedding model offline.",
        "`OLLAMA_BASE_URL` / `OLLAMA_MODEL` — local LLM endpoint and model (qwen2.5 in this project).",
        "`RAG_SEARCH_LIMIT` / `retrieval.top_k` — retrieval depth.",
        "`RAG_DISTANCE_THRESHOLD` / `distance_threshold` — score cutoff; a too-low value (2.0) historically caused hallucinations, the default is now 999.0.",
        "`CLOUD_LLM_API_KEY` / `OPENROUTER_*` — optional cloud LLM fallback (generation router).",
    ]:
        para(doc, "• " + line, size=10, space_after=2)
    cfg = reports.get("22_config_docs", {})
    keys = []
    for f in cfg.get("files", []):
        if os.path.basename(str(f.get("path", ""))) in ("config.yaml", ".env.example"):
            for c in f.get("config", []):
                if str(c) not in keys:
                    keys.append(str(c))
    if keys:
        h2(doc, "3.1 Top-level keys (config.yaml / .env.example)")
        add_field_table(doc, ["Key", "Key", "Key"],
                        [keys[i:i + 3] + [""] * max(0, 3 - len(keys[i:i + 3])) for i in range(0, len(keys), 3)],
                        widths=[2.3, 2.3, 2.3], size=8.5)
        para(doc, "The full deduplicated list of every key the code reads is in Section 6.", size=9, color=GRAY)


def section_subsystems(doc, reports):
    h1(doc, "4. Subsystem Reference")
    para(doc, "Every folder is documented below: its purpose, data flow, external dependencies, risks, and a "
              "table of the files inside it with key symbols and config keys. Subsystems are listed in the "
              "order the scan covered them (core stack first, then infrastructure, robot stack, tooling, "
              "tests, legacy, data and configuration).", space_after=10)
    for _stem, rep in ordered_reports(reports):
        render_subsystem(doc, rep)


def render_subsystem(doc, rep):
    h2(doc, rep.get("title", rep.get("folder", "Subsystem")))
    if rep.get("folder"):
        para(doc, f"Folder: `{rep['folder']}`", size=9.5, color=GRAY, space_after=4)
    para(doc, rep.get("purpose", "(no purpose captured)"))
    if rep.get("data_flow"):
        p = para(doc, None)
        set_run(p.add_run("Data flow / integration: "), size=10.5, bold=True)
        set_run(p.add_run(rep["data_flow"]), size=10.5)
    deps = rep.get("external_deps") or []
    if deps:
        para(doc, "External dependencies: " + ", ".join(str(d) for d in deps), size=10, italic=True, space_after=4)
    files = rep.get("files") or []
    if files:
        h3(doc, "Files")
        rows = []
        for f in files:
            rows.append([
                str(f.get("path", "")).replace("src/campus_helpdesk/", ""),
                str(f.get("purpose", "")).replace("\n", " "),
                ", ".join(str(s) for s in (f.get("key_symbols") or [])),
                ", ".join(str(c) for c in (f.get("config") or [])),
                str(f.get("notes", "")).replace("\n", " "),
            ])
        add_field_table(doc, ["File", "Purpose", "Key symbols", "Config", "Notes"],
                        rows, widths=[1.6, 2.2, 1.7, 1.2, 1.6])
    risks = rep.get("risks") or []
    if risks:
        p = para(doc, None)
        set_run(p.add_run("Risks / issues: "), size=10.5, bold=True)
        set_run(p.add_run("; ".join(str(r) for r in risks)), size=10.5)
    doc.add_paragraph()


def section_glossary(doc, terms):
    h1(doc, "5. Glossary")
    para(doc, "Every domain/technical term used across the codebase, deduplicated from the per-file scan. "
              "Definitions are grounded in how each term is used in this project.", space_after=8)
    if not terms:
        para(doc, "(No terms captured yet — scan reports not available.)", italic=True)
        return
    rows = [[t["term"], t["definition"]] for t in terms]
    add_field_table(doc, ["Term", "Definition"], rows, widths=[2.0, 4.7])
    para(doc, f"Total: {len(rows)} terms.", size=9, color=GRAY, space_after=0)


def section_config_reference(doc, config_keys):
    h1(doc, "6. Configuration Keys Referenced in Code")
    para(doc, "All settings keys, environment variables and config.yaml keys the code reads, deduplicated "
              "across the whole scan.", space_after=8)
    if not config_keys:
        para(doc, "(No keys captured yet — scan reports not available.)", italic=True)
        return
    add_field_table(doc, ["Key", "Key", "Key", "Key"],
                    [config_keys[i:i + 4] + [""] * max(0, 4 - len(config_keys[i:i + 4])) for i in range(0, len(config_keys), 4)],
                    widths=[1.7, 1.7, 1.7, 1.7])


def section_risks(doc, reports):
    h1(doc, "7. Known Production Risks & Notable Issues")
    para(doc, "Risks flagged in CLAUDE.md and by the code scan. These should not be silently reintroduced.", space_after=8)
    known = [
        ("Prompt injection", "`rag_chat_service.py` builds the final prompt with the raw user message while only the search query is sanitized; the sanitizer is a regex blacklist."),
        ("Streaming has no retry", "`OllamaLLMService.generate_stream` retry wrapper never executes the generator body."),
        ("VAD crash", "`vad_service.py` uses `logger` without defining it — the VAD worker crashes on voice start."),
        ("CORS / auth / XSS", "`main.py` sets `allow_origins=[\"*\"]` with `allow_credentials=True`; the API has no auth or rate limiting; the web UI renders replies via `innerHTML`."),
        ("Robot stack mocks", "`system_runtime.py` hardcodes mic `device_index=99` (always mock); `tts_service.py` Piper backend is a stub; `camera_service.py` Windows detection uses `time.asctime().startswith(\"Win\")` (always false)."),
        ("Large tracked artifacts", "The repo tracks `chunks.jsonl` (~19MB), `embedding_metadata.jsonl` (~47MB), `archive/` binaries and a nested `bvb_scraped_data/.git` — `.git` is ~900MB."),
        ("Duplicate service definitions", "There are three `llm_service.py` files (application, services, infrastructure) and two STT/TTS implementations (root + `src/services`) — historically confusing; the `src/` ones are canonical."),
        ("Analytics is dormant", "The whole `analytics/` subsystem is not published to by `application/` or `api/`."),
    ]
    for title, desc in known:
        p = doc.add_paragraph(style="List Bullet")
        set_run(p.add_run(title + ": "), size=10.5, bold=True)
        set_run(p.add_run(desc), size=10.5)
    from_scan = []
    for rep in reports.values():
        for r in rep.get("risks", []):
            r = str(r).strip()
            if r and r not in from_scan:
                from_scan.append(r)
    if from_scan:
        h2(doc, "7.1 Risks surfaced by the scan")
        for r in from_scan:
            doc.add_paragraph(r, style="List Bullet")


def section_legacy(doc, reports):
    h1(doc, "8. Legacy & Non-Production Code")
    para(doc, "Code that is not part of the production `src/` stack. Not imported by `src/`; do not extend — "
              "prefer migrating anything still needed into `src/`.", space_after=8)
    groups = [
        ("Root-level service scripts (superseded)", [
            "assistant_loop.py", "stt_service.py", "tts_service.py", "ttt_service.py",
            "presence_service.py", "benchmark.py", "chat.py", "config.py",
            "helpdesk_gui.py", "helpdesk_gui_new.py", "profile_camera.py",
        ]),
        ("Diagnostics / analysis scripts", [
            "analyze_data.py", "analyze_data_v2.py", "diag_e2e.py", "diag_ollama.py",
            "diag_rag.py", "validate_production.py", "audit_production.py",
        ]),
        ("Old top-level packages", ["services/ (root)", "conversation/ (root)"]),
        ("Old pipeline (referenced in docs)", ["bvbcet_rag_pipeline/ (removed from tree)", "archive/", "archive_tree.txt"]),
    ]
    for title, items in groups:
        para(doc, title, bold=True, size=10.5, space_after=2)
        for it in items:
            para(doc, "• " + it, size=10, space_after=2)
    legacy_rep = reports.get("20_legacy_root")
    if legacy_rep:
        render_subsystem(doc, legacy_rep)


def section_inventory(doc, reports):
    h1(doc, "9. Complete File Inventory")
    para(doc, "Every `.py` file recorded by the scan, grouped by subsystem, with line counts. Files found on disk "
              "but not recorded are listed at the end.", space_after=8)
    rows = all_file_rows(reports)
    by_sub = {}
    for path, sub, purpose in rows:
        by_sub.setdefault(sub, []).append((path, purpose))
    total_lines = 0
    for sub in sorted(by_sub):
        files = sorted(by_sub[sub])
        h2(doc, sub)
        table_rows = []
        for path, purpose in files:
            try:
                lines = sum(1 for _ in open(os.path.join(REPO_ROOT, path), encoding="utf-8", errors="replace"))
            except Exception:  # noqa: BLE001
                lines = 0
            total_lines += lines
            table_rows.append([path.replace("src/campus_helpdesk/", ""), str(lines), purpose])
        add_field_table(doc, ["File", "Lines", "Purpose"], table_rows, widths=[2.6, 0.6, 3.5], size=8.5)
    # catch files on disk not in any report
    known = {r[0] for r in rows}
    missing = []
    for py in sorted(glob.glob(os.path.join(REPO_ROOT, "**/*.py"), recursive=True)):
        rel = os.path.relpath(py, REPO_ROOT).replace("\\", "/")
        if any(part in rel for part in (".venv", "__pycache__", ".git", ".mypy_cache", ".pytest_cache", ".ruff_cache", "ffmpeg-8.1.2")):
            continue
        if rel not in known:
            missing.append(rel)
    if missing:
        h2(doc, "Files on disk not yet catalogued")
        for m in missing:
            para(doc, "• " + m, size=9.5, space_after=1)
    para(doc, f"Total lines across catalogued files: {total_lines:,}.", bold=True, space_after=0)


def section_repo_stats(doc):
    h1(doc, "10. Repository Statistics")
    stats = []
    for d in ["src", "tests", "scripts", "evaluation", "services", "conversation"]:
        root = os.path.join(REPO_ROOT, d)
        if not os.path.isdir(root):
            continue
        pys = glob.glob(os.path.join(root, "**/*.py"), recursive=True)
        pys = [p for p in pys if "__pycache__" not in p]
        lines = sum(sum(1 for _ in open(p, encoding="utf-8", errors="replace")) for p in pys)
        stats.append((d, len(pys), lines))
    total_py = sum(s[1] for s in stats)
    total_lines = sum(s[2] for s in stats)
    add_field_table(doc, ["Directory", "Python files", "Lines"],
                    [[d, str(n), f"{l:,}"] for d, n, l in stats], widths=[2.0, 1.5, 1.5])
    para(doc, f"Totals: {total_py} Python files, {total_lines:,} lines (excluding `.venv`, caches, ffmpeg and data files).",
         bold=True, space_after=0)


# ------------------------------------------------------------- main

def build_document(out_path):
    reports = load_reports()
    print(f"Loaded {len(reports)} subsystem report(s) from {SCAN_DIR}")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # footer page number
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
    for el in (fld1, instr, fld2):
        run._r.append(el)

    # ---------------- cover
    para(doc, "", space_after=0)
    para(doc, "CAMPUS HELPDESK", size=28, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, color=ACCENT)
    para(doc, "Complete Project Architecture & Codebase Reference", size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, f"Generated {datetime.date.today().isoformat()}  ·  repo root: D:\\helpdesk", size=10,
         align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY, space_after=24)
    para(doc, "Offline-first RAG chatbot for a college · Local Ollama LLM + FAISS/BM25 hybrid retrieval · "
              "Optional Raspberry Pi robot stack (STT/TTS/vision).", size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, space_after=18)
    para(doc, "This document walks through every folder and file in the repository, explains what each part is "
              "for, describes the end-to-end architecture, and defines the terminology used throughout the code.",
         size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    page_break(doc)

    # ---------------- toc
    h1(doc, "Table of Contents")
    add_toc(doc)
    page_break(doc)

    section_overview(doc)
    page_break(doc)
    section_architecture(doc)
    page_break(doc)
    section_config(doc, reports)
    page_break(doc)
    section_subsystems(doc, reports)
    page_break(doc)
    section_glossary(doc, collect_terms(reports))
    page_break(doc)
    section_config_reference(doc, collect_config(reports))
    page_break(doc)
    section_risks(doc, reports)
    page_break(doc)
    section_legacy(doc, reports)
    page_break(doc)
    section_inventory(doc, reports)
    page_break(doc)
    section_repo_stats(doc)

    doc.save(out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    out = sys.argv[1] if len(sys.argv) > 1 else os.path.join(REPO_ROOT, "Campus_Helpdesk_Complete_Architecture.docx")
    build_document(out)
